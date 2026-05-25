"""
download_utils.py — Generate downloadable TXT, PDF, and DOCX from summary text.
"""

from __future__ import annotations

import io


def to_txt(summary: str) -> bytes:
    """Return summary as UTF-8 encoded plain text bytes."""
    return summary.encode("utf-8")


def to_pdf(summary: str, title: str = "Document Summary") -> bytes:
    """Return summary as a PDF byte stream using ReportLab."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Heading1"], fontSize=16, spaceAfter=12
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=11, leading=16
    )

    # Escape HTML-special chars for ReportLab Paragraph
    safe_summary = (
        summary.replace("&", "&amp;")
               .replace("<", "&lt;")
               .replace(">", "&gt;")
               .replace("\n", "<br/>")
    )

    story = [
        Paragraph(title, title_style),
        Spacer(1, 0.3 * cm),
        Paragraph(safe_summary, body_style),
    ]
    doc.build(story)
    return buf.getvalue()


def to_docx(summary: str, title: str = "Document Summary") -> bytes:
    """Return summary as a DOCX byte stream using python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=1)
    para = doc.add_paragraph(summary)
    para.runs[0].font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
